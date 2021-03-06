# V8: Further improvement on page segmenetation
# 1. Still uses find_qn_coords as preprocessing technique, where
# the small contours that contain the question numbers ONLY are used
# to detect the question segmenetations
import cv2
import numpy as np
import pytesseract
from PIL import Image, ImageOps
from langdetect import detect
from docx import Document
from docx.shared import Inches
from GibberishDetector import classify
import os
import sys
from pdf2image import convert_from_path
import io
import re
import pandas as pd
import platform
import math as m
import ast
import shutil


def get_image(image_path):
    """Get a numpy array of an image so that one can access values[x][y]."""
    image = Image.open(image_path, 'r')
    image = image.convert('L') #makes it greyscale
    width, height = image.size
    pixel_values = list(image.getdata())
    if image.mode == 'RGB':
        channels = 3
    elif image.mode == 'L':
        channels = 1
    else:
        print("Unknown mode: %s" % image.mode)
        return None
    pixel_values = np.array(pixel_values).reshape((width, height, channels))
    return pixel_values


def is_white_image(image_name):
    numpy_array = get_image(image_name + ".jpg")
    total_pixels = numpy_array.size
    num_of_black = 0
    num_of_white = 0

    for i in numpy_array:
        for j in i:
            if j[0] > 200:
                num_of_white = num_of_white + 1
            else:
                num_of_black = num_of_black + 1

    white_percentage = num_of_white / total_pixels
    # Save as inverted image if it is a negative image
    if white_percentage < 0.75:
        # load image# Load image
        im = Image.open(image_name + ".jpg")

        # Invert
        result = ImageOps.invert(im)

        # Save
        result.save(image_name + "_inverted.jpg")
        return False
    else:
        return True


def get_thresh_and_contours(img):
    # Blurring
    imgBlur = cv2.GaussianBlur(img, (7, 7), 1)
    # cv2.imshow("blur",imgBlur)

    # convert to gray
    gray = cv2.cvtColor(imgBlur, cv2.COLOR_BGR2GRAY)
    # cv2.imshow("gray",gray)

    # threshold the grayscale image
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    # Draw contours
    result = img.copy()

    # use morphology erode to blur horizontally
    # kernel = np.ones((500,3), np.uint8)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (200, 3))  # 250,3
    morph = cv2.morphologyEx(thresh, cv2.MORPH_DILATE, kernel)

    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 17))  # previously 20,30 for eng papers
    morph = cv2.morphologyEx(morph, cv2.MORPH_OPEN, kernel)

    resized = cv2.resize(morph, (700, 850))
    # cv2.imshow("morph",resized)

    # find contours
    cntrs = cv2.findContours(morph, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cntrs = cntrs[0] if len(cntrs) == 2 else cntrs[1]
    return thresh, cntrs, result, morph


def merge_contours(thresh, cntrs, x_tolerance, y_tolerance):
    # Erase small contours, and contours which small aspect ratio (close to a square)
    for c in cntrs:
        area = cv2.contourArea(c)

        # Fill very small contours with zero (erase small contours).
        if area < 100:
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue

        # https://stackoverflow.com/questions/52247821/find-width-and-height-of-rotatedrect
        rect = cv2.minAreaRect(c)
        (x, y), (w, h), angle = rect
        #aspect_ratio = max(w, h) / min(w, h)

        # Assume zebra line must be long and narrow (long part must be at lease 1.5 times the narrow part).
        '''
        if (aspect_ratio < 1.5):
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue
            '''

    # Use "close" morphological operation to close the gaps between contours
    # https://stackoverflow.com/questions/18339988/implementing-imcloseim-se-in-opencv

    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE,
                              cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (x_tolerance, y_tolerance)));

    # Find contours in thresh_gray after closing the gaps
    cntrs, hier = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    return thresh, cntrs


def draw_contours(result, img, cntrs, image_name):
    global diagram_count
    # Contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = []
    height, width, channels = img.shape
    i = 0
    for c in cntrs:
        area = cv2.contourArea(c) / 10000
        x, y, w, h = cv2.boundingRect(c)
        cv2.rectangle(result, (x, y), (x + w, y + h), (0, 0, 255), 2)
        i += 1
        if platform.system() == "Windows":
            pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract.exe"
        
        # Read as binary image
        image = cv2.imread(image_name + ".jpg", 0)

        thresh = 255 - cv2.threshold(image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

        ROI = thresh[y:y + h, x:x + w]
        text = pytesseract.image_to_string(ROI, lang='eng', config='--psm 6')
        pseudo_text = text

        # Only add the image if it is large enough,
        # and the entire image has illegal text symbols(which is likely to be a diagram)
        #if w/width > 0.0302 and h/height > 0.0213 and y/height < 0.95:
        if w / width > 0.1 and h / height > 0.001 and y / height < 0.95 and w / h < 25:  # and (x/width < 0.4 or x/width > 0.5)
            # if is_gibberish(text) and w/h < 5:
            ############## trying out hough transform as a filter!!! ###############
            new_image = img[y:y + h, x:x + w]
            dst = cv2.Canny(new_image, 50, 200, None, 3)
            linesp = cv2.HoughLinesP(dst, 1, np.pi / 180, 50, None, 50, 2)
            if linesp is not None or is_gibberish(text):
                cv2.imwrite("TempImages/" + str(diagram_count) + ".jpg", new_image)
                document_data_list.append(("TempImages/" + str(diagram_count) + ".jpg", "image", y, pseudo_text))
                diagram_count = diagram_count + 1
            else:
                # large chunk that resembles diagram, but really is text
                document_data_list.append((text, "text", y, pseudo_text))
    return document_data_list


def is_gibberish(text):
    is_definitely_not_gibberish = False
    split = text.split("\n")

    total_value = 0
    if len(split) == 0:
        return False

    # Every line in the contour box
        for s in split:
            # if contains_quesiton_number:
            #     if there is a question number, it is definitely not gibberish. i.e we want it as a text, not image
            #     is_definitely_not_gibberish = True

            # some qn headers are auto recognized as gibberish, so we wanna make sure statements with qn numbers in them are definitely NOT gibberish as well so they are appended as text and not images
            search_sentence = re.search(r'[?0-9]+[.,]+', s, re.I)
            if search_sentence or "(a)" in s or "(b)" in s or "(c)" in s or "(d)" in s:
                is_definitely_not_gibberish = True

        gibberish_likelihood_percentage = classify(s)
        total_value = total_value + gibberish_likelihood_percentage

    if is_definitely_not_gibberish:
        return False

    average_percentage = total_value / len(split)
    if average_percentage > 45:
        # likely to be gibberish
        return True
    else:
        return False


def write_data_to_document(document_data_list, document, filename, qn_coord):
    global qn_num
    global pg_num
    global global_df
    
    # Sort data of text and images according to their y values, and add them to a word document
    document_data_list.sort(key=lambda tup: tup[2])

    for i in range(len(document_data_list)):
        data = document_data_list[i]
        item = data[0]
        typeof = data[1]
        y_coord = data[2]
        pseudo_text = data[3]
        # Eg. ("1.jpg", "image", "45", "@#!$@!$@!$")
        # STEP 1: Add qn number if current contour exceeds the page number or y_coord of the current question
        # (pg_num, y_coord, qn_num, qn_section)
        if qn_num < len(qn_coord) - 1:
            next_qn_tuple = qn_coord[qn_num + 1]
            # print((pg_num, y_coord, next_qn_tuple[0], next_qn_tuple[1]))
            if pg_num > next_qn_tuple[0] or (pg_num == next_qn_tuple[0] and y_coord > next_qn_tuple[1]):
                qn_num = qn_num + 1
                
        # STEP 2: Add question to dataframe
        if typeof == "text":
            document.add_paragraph(item)
            illegal_qn_strings = ["chij", "mark", "instructions", "go on to the next page", "blank page", "question"]
            # Do not accept text as question if it contains any of these strings
            contains_illegal_qn_string = any(ele in item.lower() for ele in illegal_qn_strings)
            if not contains_illegal_qn_string and item != "":
                # print((pg_num, y_coord, qn_num, item))
                # ['qn_num', 'pg_num', 'pdf_name', 'text', 'image_path']
                if qn_num not in global_df.index:
                    global_df.loc[qn_num] = [qn_num, pg_num, filename, item, ""]
                else:
                    global_df.loc[qn_num] = [qn_num, pg_num, filename, global_df.loc[qn_num][3] + item, global_df.loc[qn_num][4]]
            
        elif typeof == "image":
            if qn_num not in global_df.index:
                    global_df.loc[qn_num] = [qn_num, pg_num, filename, "", item]
            else:
                global_df.loc[qn_num] = [qn_num, pg_num, filename, global_df.loc[qn_num][3], global_df.loc[qn_num][4] + ";" + item]
            document.add_picture(item, width=Inches(5))
        


def generate_document(filename, documentdir, qn_coord):
    global pg_num
    global total_pages

    print("Step 2 (Output Generation): PG " + str(pg_num) + "/" + str(total_pages))
    image_name = filename.replace(".jpg", "")
    document = Document()

    ###### Step 1: Convert to positive image if image is negative
    if not is_white_image(image_name):
        image_name = image_name + "_inverted"

    ###### Step 2: Get the initial thresh and contours
    img = cv2.imread(image_name + ".jpg")
    height, width, channels = img.shape
    thresh, cntrs, result, morph = get_thresh_and_contours(img)

    ###### Step 3: Merge contours that are close together
    # Modify the x and y tolerance to change how far it must be before it will merge!
    x_tolerance = m.floor(0.18 * width) # previously 300px
    y_tolerance = m.floor(0.009 * height) # previously 0.014964 #SX: 0.015
    thresh, cntrs = merge_contours(thresh, cntrs, x_tolerance, y_tolerance)

    ###### Step 4: Draw the contours on the image
    # ordered_value_tuples contains ordered tuples of (text, y_coord)
    # document_data_list contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = draw_contours(result, img, cntrs, image_name)
    # cv2.imwrite("contour_img/" + str(file_count) + ".jpg", result)

    ###### Step 5: Write and Save to a new Microsoft Word Document
    write_data_to_document(document_data_list, document, filename, qn_coord)
    # Remove /images from image_name. example image_name is images/P6_English_2019_CA1_CHIJ/pg_1_P6_English_2019_CA1_CHIJ.jpg
    image_name = image_name.split('/', 1)[1]
    # Test paper name found in /images, example parentdir is P6_English_2019_CA1_CHIJ
    parentdir = image_name.split('/', 1)[0]

    # put this under documentdir
    if not os.path.exists(documentdir + "/" + parentdir):
        os.makedirs(documentdir + "/" + parentdir)

    document.save(documentdir + "/" + image_name + ".docx")

    # cv2.imshow("THRESH", thresh)
    # cv2.imshow("MORPH", morph)

    ####### Step 6: Display results
    ims=cv2.resize(result,(700,850))
    cv2.imwrite("TempContours/" + str(pg_num) + ".jpg", ims)
    pg_num = pg_num + 1
    # cv2.imshow("RESULT", ims)
    # cv2.waitKey(0)
    # cv2.destroyAllWindows()

# Copies all files from src directory to dest directory
def copytree(src, dst, symlinks=False, ignore=None):
    if not os.path.exists(dst):
        os.makedirs(dst)
    for item in os.listdir(src):
        s = os.path.join(src, item)
        d = os.path.join(dst, item)
        if os.path.isdir(s):
            copytree(s, d, symlinks, ignore)
        else:
            if not os.path.exists(d) or os.stat(s).st_mtime - os.stat(d).st_mtime > 1:
                shutil.copy2(s, d)
                
# Map the page number and y coordinates of each question     
def find_qn_coords(filenames_list):
    global total_pages
    qn_coord = []
    qn_coord.append((0, 0))
    pg_num = 1
    qn_num = 1
    diagram_count = 1

    for filename in filenames_list:
        print("Step 1 (Preprocessing): PG " + str(pg_num) + "/" + str(total_pages))
        # if pg_num < 24:
        #     continue

        image_name = filename.replace(".jpg", "")
        ###### Step 1: Convert to positive image if image is negative
        if not is_white_image(image_name):
            image_name = image_name + "_inverted"

        ###### Step 2: Get the initial thresh and contours
        img = cv2.imread(image_name + ".jpg")
        height, width, channels = img.shape
        thresh, cntrs, result, morph = get_thresh_and_contours(img)

        ###### Step 3: Merge contours that are close together
        # Modify the x and y tolerance to change how far it must be before it will merge!
        x_tolerance = m.floor(0.02138*width) # previously 300px
        y_tolerance = m.floor(0.024964*height) # previously 35px
        thresh, cntrs = merge_contours(thresh, cntrs, x_tolerance, y_tolerance)
        
        sorted_cntr_tuples = []
        for c in cntrs:
            area = cv2.contourArea(c) / 10000
            x, y, w, h = cv2.boundingRect(c)
            cv2.rectangle(result, (x, y), (x + w, y + h), (0, 0, 255), 2)

            if area < 0.1 and area > 0.01 and y/height < 0.855 and x/width < 0.25 and w/h < 2 and w/h > 0.5:
                new_image = img[y:y + h, x:x + w]
                text = pytesseract.image_to_string(new_image, lang='eng', config='--psm 6')
                if text != "":
                    # Check if pseudo_text contains numbers contained in any brackets
                    #matches = re.search(r'[\[\(\|\{][0-9a-z][\]\)\}\|]', text, re.I)
                    illegal_qn_strings = ["(", ")", "[", "]", "{", "}", "|", "NO"]
                    # Do not accept text as question if it contains any of these strings
                    contains_illegal_qn_string = any(ele in text.lower() for ele in illegal_qn_strings)

                    if not contains_illegal_qn_string:
                        sorted_cntr_tuples.append((c, pg_num, y, w, h, x/width))

        sorted_cntr_tuples.sort(key=lambda tup: tup[2])
        # Comment this out to visualise the processing of small contours under TempImages/small_[NAME].jpg
        small_cntrs = []
        for c, pg_num, y, w, h, xw in sorted_cntr_tuples:
            x, y, w, h = cv2.boundingRect(c)
            new_image = img[y:y + h, x:x + w]
            cv2.imwrite("TempImages/small_" + str(diagram_count) + ".jpg", new_image)
            # Read as binary image
            small_image = cv2.imread("TempImages/small_" + str(diagram_count) + ".jpg", 0)
            text = pytesseract.image_to_string(small_image, lang='eng', config='--psm 6')
            small_cntrs.append((pg_num, y))
            diagram_count = diagram_count + 1

        for c, pg_num, y, w, h, xw in sorted_cntr_tuples:
            qn_coord.append((pg_num, y))
            # print((pg_num, y))

        cv2.imwrite("TempContours/" + str(pg_num) + ".jpg", result)
        pg_num = pg_num + 1
    
    return qn_coord
    


def main(pdfname):
    global total_pages
    global global_df

    print(pdfname)
    paper_name = pdfname.replace(".pdf", "")
    pdf_path = "Sample Resources/" + paper_name + ".pdf"
    pages = convert_from_path(pdf_path)
    pg_cntr = 1
    filenames_list = []

    sub_dir = str("images/" + pdf_path.split('/')[-1].replace('.pdf', '') + "/")
    if not os.path.exists(sub_dir):
        os.makedirs(sub_dir)

    for page in pages:
        filename = "pg_" + str(pg_cntr) + '_' + pdf_path.split('/')[-1].replace('.pdf', '.jpg')
        page.save(sub_dir + filename)
        pg_cntr = pg_cntr + 1
        filenames_list.append(sub_dir + filename)

    total_pages = len(filenames_list)
    qn_coord = find_qn_coords(filenames_list)
    # qn_coord = ast.literal_eval("[(0, 0, 0, 0), (2, 1365, 1, ''), (2, 1703, 2, ''), (3, 1131, 3, ''), (3, 1819, 4, ''), (4, 966, 5, ''), (4, 1779, 6, ''), (5, 2056, 7, ''), (6, 744, 8, ''), (6, 1934, 9, ''), (7, 757, 10, ''), (7, 1924, 11, ''), (8, 905, 12, ''), (8, 1710, 13, ''), (9, 1077, 14, ''), (9, 1914, 15, ''), (10, 1256, 16, ''), (11, 1630, 17, ''), (12, 1385, 18, ''), (13, 1432, 19, ''), (14, 1401, 20, ''), (15, 1292, 21, ''), (16, 1047, 22, ''), (17, 1295, 23, ''), (18, 1169, 24, ''), (19, 1005, 25, ''), (19, 1527, 26, ''), (20, 1535, 27, ''), (21, 1186, 28, ''), (21, 1968, 29, ''), (24, 1630, 30, 2), (26, 1591, 31, 2), (27, 1584, 32, 2), (29, 268, 33, 3), (30, 1935, 34, 2), (31, 1858, 35, 3), (32, 1035, 36, 3), (33, 1711, 37, 1), (34, 1553, 38, 1), (35, 1395, 39, 1), (36, 1739, 40, 3)]")

    for filename in filenames_list:
        generate_document(filename, "OutputDocuments4", qn_coord)

    global_df.to_csv("output.csv")

    # Copies all the output to a new folder under Output/PDF NAME
    dirpath = os.getcwd()
    copytree(dirpath + "/TempContours", dirpath + "/Output/" + paper_name + "/TempContours")
    copytree(dirpath + "/TempImages", dirpath + "/Output/" + paper_name + "/TempImages")
    copytree(dirpath + "/images/" + paper_name, dirpath + "/Output/" + paper_name + "/images")
    shutil.copyfile(dirpath + "/output.csv", dirpath + "/Output/" + paper_name + "/output.csv")


qn_num = 1
pg_num = 1
diagram_count = 1
total_pages = -1
global_df = pd.DataFrame(columns=['qn_num', 'pg_num', 'pdf_name', 'text', "img_path"])

for filename in os.listdir("Sample Resources"):
    if filename.endswith(".pdf"): 
        main(filename)