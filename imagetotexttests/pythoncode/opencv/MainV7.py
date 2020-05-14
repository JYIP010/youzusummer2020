# V7: Current Implementation

import cv2
import numpy as np
import pytesseract
from PIL import Image, ImageOps
from langdetect import detect
from docx import Document
from docx.shared import Inches
from GibberishDetector import classify
import os
import sys
from pdf2image import convert_from_path
import io
import re
import pandas as pd
import platform
import math as m


def get_image(image_path):
    """Get a numpy array of an image so that one can access values[x][y]."""
    image = Image.open(image_path, 'r')
    width, height = image.size
    pixel_values = list(image.getdata())
    if image.mode == 'RGB':
        channels = 3
    elif image.mode == 'L':
        channels = 1
    else:
        print("Unknown mode: %s" % image.mode)
        return None
    pixel_values = np.array(pixel_values).reshape((width, height, channels))
    return pixel_values


def is_white_image(image_name):
    numpy_array = get_image(image_name + ".jpg")
    total_pixels = numpy_array.size
    num_of_white = np.count_nonzero(numpy_array == [255, 255, 255])
    num_of_black = np.count_nonzero(numpy_array == [1, 1, 1])
    white_percentage = num_of_white / total_pixels
    # Save as inverted image if it is a negative image
    if white_percentage < 0.75:
        # load image# Load image
        im = Image.open(image_name + ".jpg")

        # Invert
        result = ImageOps.invert(im)

        # Save
        result.save(image_name + "_inverted.jpg")
        return False
    else:
        return True


def get_thresh_and_contours(img):
    # Blurring
    imgBlur = cv2.GaussianBlur(img, (7, 7), 1)
    # cv2.imshow("blur",imgBlur)

    # convert to gray
    gray = cv2.cvtColor(imgBlur, cv2.COLOR_BGR2GRAY)
    # cv2.imshow("gray",gray)

    # threshold the grayscale image
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    # Draw contours
    result = img.copy()

    # use morphology erode to blur horizontally
    # kernel = np.ones((500,3), np.uint8)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (200, 3))  # 250,3
    morph = cv2.morphologyEx(thresh, cv2.MORPH_DILATE, kernel)

    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (20, 30))  # previously 3,17
    morph = cv2.morphologyEx(morph, cv2.MORPH_OPEN, kernel)

    resized = cv2.resize(morph, (700, 850))
    # cv2.imshow("morph",resized)

    # find contours
    cntrs = cv2.findContours(morph, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cntrs = cntrs[0] if len(cntrs) == 2 else cntrs[1]
    return thresh, cntrs, result, morph


def merge_contours(thresh, cntrs, x_tolerance, y_tolerance):
    # Erase small contours, and contours which small aspect ratio (close to a square)
    for c in cntrs:
        area = cv2.contourArea(c)

        # Fill very small contours with zero (erase small contours).
        if area < 100:
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue

        # https://stackoverflow.com/questions/52247821/find-width-and-height-of-rotatedrect
        rect = cv2.minAreaRect(c)
        (x, y), (w, h), angle = rect
        #aspect_ratio = max(w, h) / min(w, h)

        # Assume zebra line must be long and narrow (long part must be at lease 1.5 times the narrow part).
        '''
        if (aspect_ratio < 1.5):
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue
            '''

    # Use "close" morphological operation to close the gaps between contours
    # https://stackoverflow.com/questions/18339988/implementing-imcloseim-se-in-opencv

    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE,
                              cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (x_tolerance, y_tolerance)));

    # Find contours in thresh_gray after closing the gaps
    cntrs, hier = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    return thresh, cntrs


def draw_contours(result, img, cntrs, image_name):
    global diagram_count
    # Contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = []
    height, width, channels = img.shape
    i = 0
    for c in cntrs:
        area = cv2.contourArea(c) / 10000
        x, y, w, h = cv2.boundingRect(c)
        cv2.rectangle(result, (x, y), (x + w, y + h), (0, 0, 255), 2)
        i += 1
        if platform.system() == "Windows":
            pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract.exe"
        
        # Read as binary image
        image = cv2.imread(image_name + ".jpg", 0)

        thresh = 255 - cv2.threshold(image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

        ROI = thresh[y:y + h, x:x + w]
        text = pytesseract.image_to_string(ROI, lang='eng', config='--psm 6')
        pseudo_text = text

        # Only add the image if it is large enough,
        # and the entire image has illegal text symbols(which is likely to be a diagram)
        #if w/width > 0.0302 and h/height > 0.0213 and y/height < 0.95:
        if w/width > 0.1 and h/height > 0.003 and y/height < 0.95: #and (x/width < 0.4 or x/width > 0.5)
            #if is_gibberish(text) and w/h < 5:
            if is_gibberish(text) or 0.35 < (w*h)/(width*height) < 0.97:
                if h/height > 0.1 and w/h < 10:
                    # Likely to be an image
                    new_image = img[y:y + h, x:x + w]
                    cv2.imwrite("TempImages/" + str(diagram_count) + ".jpg", new_image)
                    document_data_list.append(("TempImages/" + str(diagram_count) + ".jpg", "image", y, pseudo_text))
                    diagram_count = diagram_count + 1
                else:
                    # Likely to be text, just small regions like "Go on to the next page"
                    document_data_list.append((text, "text", y, pseudo_text))
            else:
                # Likely to be a text
                document_data_list.append((text, "text", y, pseudo_text))

    return document_data_list


def is_gibberish(text):
    is_definitely_not_gibberish = False
    split = text.split("\n")

    total_value = 0
    if len(split) == 0:
        return False

    # Every line in the contour box
    for s in split:
        # if contains_quesiton_number:
        #     if there is a question number, it is definitely not gibberish. i.e we want it as a text, not image
        #     is_definitely_not_gibberish = True
        if "(1)" in s or "(2)" in s or "(3)" in s or "(4)" in s or "(a)" in s or "(b)" in s or "(c)" in s or "(d)" in s:
            # if there is a answer number, it is definitely not gibberish. i.e we want it as a text, not image
            is_definitely_not_gibberish = True

        gibberish_likelihood_percentage = classify(s)
        total_value = total_value + gibberish_likelihood_percentage

    if is_definitely_not_gibberish:
        return False

    average_percentage = total_value / len(split)
    if average_percentage > 30:
        # likely to be gibberish
        return True
    else:
        return False


def write_data_to_document(document_data_list, document, filename):
    global qn_num
    global pg_num
    global global_df
    
    # Sort data of text and images according to their y values, and add them to a word document
    document_data_list.sort(key=lambda tup: tup[2])
    can_add_qn_num = False
    
    for i in range(len(document_data_list)):
        data = document_data_list[i]
        #data_list=list(data)
        #data_list.insert(0,qn_num)
        
        item = data[0]
        typeof = data[1]
        y_coord = data[2]
        pseudo_text = data[3]
        # Eg. ("1.jpg", "image", "45", "@#!$@!$@!$")
        
        # STEP 1: Add question to dataframe
        if typeof == "text":
            document.add_paragraph(item)
            illegal_qn_strings = ["chij", "mark", "instructions", "go on to the next page", "blank page", "question"]
            # Do not accept text as question if it contains any of these strings
            contains_illegal_qn_string = any(ele in item.lower() for ele in illegal_qn_strings)
            if not contains_illegal_qn_string:
                print((qn_num, item))
                # ['qn_num', 'pg_num', 'pdf_name', 'text', 'image_path']
                if qn_num not in global_df.index:
                    global_df.loc[qn_num] = [qn_num, pg_num, filename, item, ""]
                else:
                    global_df.loc[qn_num] = [qn_num, pg_num, filename, global_df.loc[qn_num][3] + item, global_df.loc[qn_num][4]]
            
        elif typeof == "image":
            if qn_num not in global_df.index:
                    global_df.loc[qn_num] = [qn_num, pg_num, filename, "", item]
            else:
                global_df.loc[qn_num] = [qn_num, pg_num, filename, global_df.loc[qn_num][3], global_df.loc[qn_num][4] + ";" + item]
            document.add_picture(item, width=Inches(5))

        # STEP 2: Check if qn_num should be added in the NEXT contour
        already_added_qn_num = False
        qn_list = ["(1)", "(2)", "(3)", "(4)"]
        # Check if text contains any string found in qn_list
        contains_string = any(ele in pseudo_text for ele in qn_list)
        if contains_string:
            can_add_qn_num = True
            already_added_qn_num = True
                
        if "?" in pseudo_text:
            # If there is a ? symbol in the current box,
            # Ignore the (1), (2), (3), (4) ans num in the next box
            if i == len(document_data_list) - 1:
                # no more bounding boxes below in the same page
                can_add_qn_num = True
                already_added_qn_num = True
            else:
                next_data = document_data_list[i + 1] 
                next_data_contains_string = any(ele in next_data[3] for ele in qn_list)
                if next_data_contains_string:
                    #if next data below contains answer options, do not add qn number yet
                    pass
                else:
                    can_add_qn_num = True
                    already_added_qn_num = True

        if not already_added_qn_num:
            # Check regex for square brackets. i.e [2m]
            search_string = re.search(r'[\[\(\|\{]+[1-9]+[o0-9]*[mn]*[\]\)\}\|]+', pseudo_text, re.I)
            if search_string:
                can_add_qn_num = True
                already_added_qn_num = True
                
        # STEP 3: Check if the qn_num should be increased in the CURRENT contour
        # Only allow this if there is no diagram in the NEXT contour
        if i == len(document_data_list) - 1 and can_add_qn_num:
            # no more bounding boxes below in the same page
            can_add_qn_num = False
            qn_num = qn_num + 1
        elif i != len(document_data_list) - 1 and can_add_qn_num:
            next_data = document_data_list[i + 1] 
            if next_data[1] == "text":
                can_add_qn_num = False
                qn_num = qn_num + 1
                
    pg_num = pg_num + 1
        


def generate_document(filename, documentdir):
    image_name = filename.replace(".jpg", "")
    document = Document()

    ###### Step 1: Convert to positive image if image is negative
    if not is_white_image(image_name):
        image_name = image_name + "_inverted"

    ###### Step 2: Get the initial thresh and contours
    img = cv2.imread(image_name + ".jpg")
    height, width, channels = img.shape
    thresh, cntrs, result, morph = get_thresh_and_contours(img)

    ###### Step 3: Merge contours that are close together
    # Modify the x and y tolerance to change how far it must be before it will merge!
    x_tolerance = m.floor(0.18138*width) # previously 300px
    y_tolerance = m.floor(0.014964*height) # previously 35px
    thresh, cntrs = merge_contours(thresh, cntrs, x_tolerance, y_tolerance)

    ###### Step 4: Draw the contours on the image
    # ordered_value_tuples contains ordered tuples of (text, y_coord)

    # document_data_list contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = draw_contours(result, img, cntrs, image_name)

    ###### Step 5: Write and Save to a new Microsoft Word Document
    write_data_to_document(document_data_list, document, filename)
    # Remove /images from image_name. example image_name is images/P6_English_2019_CA1_CHIJ/pg_1_P6_English_2019_CA1_CHIJ.jpg
    image_name = image_name.split('/', 1)[1]
    # Test paper name found in /images, example parentdir is P6_English_2019_CA1_CHIJ
    parentdir = image_name.split('/', 1)[0]

    # put this under documentdir
    if not os.path.exists(documentdir + "/" + parentdir):
        os.makedirs(documentdir + "/" + parentdir)

    document.save(documentdir + "/" + image_name + ".docx")

    # cv2.imshow("THRESH", thresh)
    # cv2.imshow("MORPH", morph)

    ###### Step 6: Display results
    # ims=cv2.resize(result,(700,850))
    # cv2.imshow("RESULT", ims)
    # cv2.waitKey(0)
    # cv2.destroyAllWindows()


# for filename in os.listdir("Sample Resources"):
#     if filename.endswith(".jpg") and not filename.endswith("_inverted.jpg"):
#         print(filename)
#         generate_document(filename, "OutputDocuments4")
#         pass

pdf_path = "Sample Resources/P6_English_2019_CA1_CHIJ.pdf"
pages = convert_from_path(pdf_path)
pg_cntr = 1
filenames_list = []

sub_dir = str("images/" + pdf_path.split('/')[-1].replace('.pdf', '') + "/")
if not os.path.exists(sub_dir):
    os.makedirs(sub_dir)

for page in pages:
    filename = "pg_" + str(pg_cntr) + '_' + pdf_path.split('/')[-1].replace('.pdf', '.jpg')
    page.save(sub_dir + filename)
    pg_cntr = pg_cntr + 1
    filenames_list.append(sub_dir + filename)

qn_num = 1
pg_num = 1
diagram_count = 1
global_df = pd.DataFrame(columns=['qn_num', 'pg_num', 'pdf_name', 'text', "img_path"])

for file in filenames_list:
    generate_document(file, "OutputDocuments4")
    
global_df.to_csv("output.csv")