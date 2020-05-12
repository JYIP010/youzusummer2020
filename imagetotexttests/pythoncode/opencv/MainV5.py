# V5: Current implementation

import cv2
import numpy as np
import pytesseract
from PIL import Image, ImageOps
from langdetect import detect
from docx import Document
from docx.shared import Inches
from GibberishDetector import classify
import os 

def get_image(image_path):
    """Get a numpy array of an image so that one can access values[x][y]."""
    image = Image.open(image_path, 'r')
    width, height = image.size
    pixel_values = list(image.getdata())
    if image.mode == 'RGB':
        channels = 3
    elif image.mode == 'L':
        channels = 1
    else:
        print("Unknown mode: %s" % image.mode)
        return None
    pixel_values = np.array(pixel_values).reshape((width, height, channels))
    return pixel_values

def is_white_image(image_name):
    numpy_array = get_image("Sample Resources/" + image_name + ".jpg")
    total_pixels = numpy_array.size
    num_of_white = np.count_nonzero(numpy_array == [255,255, 255])
    num_of_black = np.count_nonzero(numpy_array == [1,1, 1])
    white_percentage = num_of_white / total_pixels
    #Save as inverted image if it is a negative image
    if white_percentage < 0.75:
        # load image# Load image 
        im = Image.open("Sample Resources/" + image_name + ".jpg")

        # Invert
        result = ImageOps.invert(im)

        # Save
        result.save("Sample Resources/" + image_name + "_inverted.jpg")
        return False
    else:
        return True

def get_thresh_and_contours(img):
    #Blurring
    imgBlur = cv2.GaussianBlur(img, (7, 7), 1)
    # cv2.imshow("blur",imgBlur)

    # convert to gray
    gray = cv2.cvtColor(imgBlur, cv2.COLOR_BGR2GRAY)
    # cv2.imshow("gray",gray)

    # threshold the grayscale image
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    # Draw contours
    result = img.copy()

    # use morphology erode to blur horizontally
    #kernel = np.ones((500,3), np.uint8)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (200, 3)) #250,3
    morph = cv2.morphologyEx(thresh, cv2.MORPH_DILATE, kernel)

    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 17)) #3,17
    morph = cv2.morphologyEx(morph, cv2.MORPH_OPEN, kernel)

    resized=cv2.resize(morph,(700,850))
    #cv2.imshow("morph",resized)

    # find contours
    cntrs = cv2.findContours(morph, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cntrs = cntrs[0] if len(cntrs) == 2 else cntrs[1]
    return thresh, cntrs, result, morph

def merge_contours(thresh, cntrs, x_tolerance, y_tolerance):
    # Erase small contours, and contours which small aspect ratio (close to a square)
    for c in cntrs:
        area = cv2.contourArea(c)

        # Fill very small contours with zero (erase small contours).
        if area < 100:
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue

        # https://stackoverflow.com/questions/52247821/find-width-and-height-of-rotatedrect
        rect = cv2.minAreaRect(c)
        (x, y), (w, h), angle = rect
        aspect_ratio = max(w, h) / min(w, h)

        # Assume zebra line must be long and narrow (long part must be at lease 1.5 times the narrow part).
        if (aspect_ratio < 1.5):
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue


    # Use "close" morphological operation to close the gaps between contours
    # https://stackoverflow.com/questions/18339988/implementing-imcloseim-se-in-opencv

    
    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (x_tolerance,y_tolerance)));

    # Find contours in thresh_gray after closing the gaps
    cntrs, hier = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    return thresh, cntrs

def draw_contours(result, img, cntrs, image_name):
    # Contains ordered tuples of (text, y_coord)
    ordered_value_tuples = []

    # Contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = []

    i=0
    count = 0
    for c in cntrs:
        area = cv2.contourArea(c)/10000
        x,y,w,h = cv2.boundingRect(c)
        # if h < 50 and w > 400 :
        if True:
            cv2.rectangle(result, (x, y), (x+w, y+h), (0, 0, 255), 2)
            # print("Box: " + str(i) + ": (" + str(int(x)) + ", " + str(int(y)) + "," + str(int(w)) + "," + str(int(h)) + ")")
            # print('Area: ' + str(area))
            i += 1
            pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract.exe"

            image = cv2.imread("Sample Resources/" + image_name + ".jpg", 0)
            thresh = 255 - cv2.threshold(image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

            ROI = thresh[y:y+h,x:x+w]
            text = pytesseract.image_to_string(ROI, lang='eng',config='--psm 6')
            
            # Only add the image if it is large enough, 
            # and the entire image has illegal text symbols(which is likely to be a diagram)
            if w > 50 and h > 50:
                if is_gibberish(text):
                    new_image = img[y:y+h, x:x+w]
                    cv2.imwrite("TempImages/" + str(count) + ".jpg" , new_image)
                    document_data_list.append(("TempImages/" + str(count) + ".jpg", "image", y))
                    count = count + 1

            ordered_value_tuples.append((text, y))
            ordered_value_tuples.sort(key = lambda tup: tup[1])

    for value in ordered_value_tuples:
        text = value[0]
        y_coord = value[1]

        if not is_gibberish(text):
            document_data_list.append((text, "text", y_coord))

    return document_data_list

def is_gibberish(text):
    is_definitely_not_gibberish = False
    split = text.split("\n")
    
    total_value = 0
    if len(split) == 0:
        return False

    for s in split:
        # if contains_quesiton_number:
        #     if there is a question number, it is definitely not gibberish. i.e we want it as a text, not image
        #     is_definitely_not_gibberish = True
        if "(1)" in s or "(2)" in s or "(3)" in s or "(4)" in s or "(a)" in s or "(b)" in s or "(c)" in s or "(d)" in s:
            # if there is a answer number, it is definitely not gibberish. i.e we want it as a text, not image
            is_definitely_not_gibberish = True

        gibberish_likelihood_percentage = classify(s)
        total_value = total_value + gibberish_likelihood_percentage
    
    if is_definitely_not_gibberish:
        return False

    average_percentage = total_value / len(split)
    print(text)
    print(average_percentage)
    print("--------------")
    if average_percentage > 25:
        #likely to be gibberish
        return True
    else:
        return False

def write_data_to_document(document_data_list, document):
    # Sort data of text and images according to their y values, and add them to a word document
    document_data_list.sort(key = lambda tup: tup[2])
    for data in document_data_list:
        item = data[0]
        typeof = data[1]
        y_coord = data[2]
        if typeof == "text":
            document.add_paragraph(item)
        elif typeof == "image":
            document.add_picture(item, width=Inches(5))

def generate_document(filename, documentdir):
    image_name = filename.replace(".jpg", "")
    document = Document()

    ###### Step 1: Convert to positive image if image is negative
    if not is_white_image(image_name):
        image_name = image_name + "_inverted"

    ###### Step 2: Get the initial thresh and contours
    img = cv2.imread("Sample Resources/" + image_name + ".jpg")
    thresh, cntrs, result, morph = get_thresh_and_contours(img)
    
    ###### Step 3: Merge contours that are close together 
    # Modify the x and y tolerance to change how far it must be before it will merge!
    x_tolerance = 300
    y_tolerance = 35
    thresh, cntrs = merge_contours(thresh, cntrs, x_tolerance, y_tolerance)

    
    ###### Step 4: Draw the contours on the image 
    # ordered_value_tuples contains ordered tuples of (text, y_coord)

    # document_data_list contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = draw_contours(result, img, cntrs, image_name)

    ###### Step 5: Write and Save to a new Microsoft Word Document 
    write_data_to_document(document_data_list, document)
    document.save(documentdir + "/" + image_name + ".docx")

    # cv2.imshow("THRESH", thresh)
    # cv2.imshow("MORPH", morph)

    ###### Step 6: Display results 
    ims=cv2.resize(result,(700,850))
    cv2.imshow("RESULT", ims)
    cv2.waitKey(0)
    cv2.destroyAllWindows()


for filename in os.listdir("Sample Resources"):
    if filename.endswith(".jpg") and not filename.endswith("_inverted.jpg"): 
        #print(filename)
        #generate_document(filename, "OutputDocuments4")
        pass

generate_document("pg_2_P6_Science_2019_SA2_CHIJ", "OutputDocuments4")